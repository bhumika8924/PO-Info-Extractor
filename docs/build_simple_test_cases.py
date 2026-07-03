from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "PO_Info_Extractor_Manual_Test_Cases.docx"
OUTPUT = HERE / "PO_Info_Extractor_Manual_Test_Cases.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(35, 46, 58)
MUTED = RGBColor(102, 112, 133)

SECTION_STARTS = {
    1: "Startup and Navigation",
    6: "Manual PDF Upload",
    20: "Extraction Results",
    28: "Auto Upload",
    35: "History and Downloads",
    43: "Reliability, Security, and User Interface",
}


def set_font(run, size, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def extract_cases(source_path):
    source = Document(source_path)
    cases = []
    for table in source.tables:
        if not table.rows:
            continue
        header = table.rows[0].cells[0].text.strip()
        match = re.match(r"TC-(\d{2})\s*\|\s*(.+)", header)
        if not match or len(table.rows) < 5:
            continue
        number = int(match.group(1))
        title = match.group(2).strip()
        procedure = table.rows[3].cells[1].text.strip()
        expected = table.rows[4].cells[1].text.strip()
        cases.append((number, title, procedure, expected))
    if len(cases) != 50:
        raise RuntimeError(f"Expected 50 cases, found {len(cases)}")
    return sorted(cases)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = DARK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = BLUE
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.clear()
    run = header.add_run("PO INFO EXTRACTOR  |  MANUAL TEST CASES")
    set_font(run, 8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)


def add_case(doc, number, title, procedure, expected):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(13)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(f"{number:02d} {title}")
    set_font(run, 15, bold=True, color=DARK)

    procedure_p = doc.add_paragraph()
    procedure_p.paragraph_format.left_indent = Inches(0.08)
    procedure_p.paragraph_format.space_after = Pt(8)
    procedure_p.paragraph_format.line_spacing = 1.25
    procedure_p.paragraph_format.keep_with_next = True
    procedure_run = procedure_p.add_run(procedure)
    set_font(procedure_run, 11.5, color=DARK)

    expected_p = doc.add_paragraph()
    expected_p.paragraph_format.left_indent = Inches(0.08)
    expected_p.paragraph_format.space_after = Pt(30)
    expected_p.paragraph_format.line_spacing = 1.25
    expected_run = expected_p.add_run(expected)
    set_font(expected_run, 11.5, color=DARK)


def build():
    cases = extract_cases(SOURCE)
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("PO Info Extractor Manual Test Cases")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle.add_run("Procedure and expected result reference")
    set_font(subtitle_run, 12.5, italic=True, color=MUTED)

    section_case_count = 0
    for number, case_title, procedure, expected in cases:
        if number in SECTION_STARTS:
            if number != 1:
                doc.add_page_break()
            section_heading = doc.add_paragraph(style="Heading 1")
            section_heading.add_run(SECTION_STARTS[number])
            section_case_count = 0

        add_case(doc, number, case_title, procedure, expected)
        section_case_count += 1
        remaining_in_section = min(
            [start for start in SECTION_STARTS if start > number] or [51]
        ) - number - 1
        if section_case_count % 2 == 0 and remaining_in_section > 0:
            doc.add_page_break()

    properties = doc.core_properties
    properties.title = "PO Info Extractor Manual Test Cases"
    properties.subject = "Simplified manual test cases with procedure and expected result"
    properties.author = "PO Info Extractor QA"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
